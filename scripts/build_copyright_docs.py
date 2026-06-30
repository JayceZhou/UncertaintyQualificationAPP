"""Build copyright application and product manual from the supplied templates."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docs"
ASSETS = ROOT / "docs" / "assets"
APPLICATION_TEMPLATE = Path(
    "/Users/jayce/Documents/博士/专利/软著初稿-上海大学-钱老师-激光焊接全流程数据库管理系统/"
    "初稿-上大-钱老师-1-申请表-20260507.docx"
)
APP_NAME = "基于变分证据采样的异构材料数据不确定性分析系统"
VERSION = "V1.0"


def core_source_line_count() -> int:
    files = [ROOT / "app.py"]
    for package in ("appcore", "ui", "uqcore", "inference"):
        files.extend(sorted((ROOT / package).glob("*.py")))
    return sum(len(path.read_text(encoding="utf-8").splitlines()) for path in files)


def set_cn_font(run, name: str, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_cn_font(run, "Arial", 9)


def add_bottom_border(paragraph, color="808080", size="6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(0)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in (
        ("Heading 1", 14),
        ("Heading 2", 12),
        ("Heading 3", 12),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    table.autofit = False
    table.columns[0].width = Cm(14.3)
    table.columns[1].width = Cm(1.7)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cn_font(left.add_run(f"{APP_NAME} {VERSION}"), "SimSun", 9)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(right)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 0, 0, 0, 0)
    add_bottom_border(left)
    add_bottom_border(right)


def add_body(document: Document, text: str, first_indent=True) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_indent:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    set_cn_font(paragraph.add_run(text), "SimSun", 10.5)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.first_line_indent = Cm(-0.35)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        set_cn_font(paragraph.add_run(item), "SimSun", 10.5)


def add_steps(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Cm(0.8)
        paragraph.paragraph_format.first_line_indent = Cm(-0.4)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        set_cn_font(paragraph.add_run(item), "SimSun", 10.5)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Cm(widths[index])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.5
        set_cn_font(paragraph.add_run(header), "SimSun", 10.5, bold=True)
        set_cell_margins(cell)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].width = Cm(widths[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == len(headers) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.5
            set_cn_font(paragraph.add_run(value), "SimSun", 10.5)
            set_cell_margins(cells[index])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_screenshot(document: Document, filename: str, caption: str) -> None:
    path = ASSETS / filename
    if not path.is_file():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Cm(15.0))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.line_spacing = 1.5
    caption_paragraph.paragraph_format.space_after = Pt(0)
    set_cn_font(caption_paragraph.add_run(caption), "SimSun", 10.5)


def page_break(document: Document) -> None:
    # The source template uses natural page flow. Explicit break paragraphs can
    # be pushed to a new page by a preceding table and create blank pages.
    return None


def build_manual() -> Path:
    document = Document()
    configure_document(document)

    document.add_heading("一、软件概述", level=1)
    document.add_heading("1.1 软件简介", level=2)
    add_body(
        document,
        "本系统面向粉末衍射分类结果和二维材料矢量场预测结果，提供统一的不确定性量化、风险分级、"
        "高风险对象筛选与可视化分析能力。系统将 DFUN 工作中的 Monte Carlo Dropout 多次随机前向"
        "统计方法，与 EviField 工作中的二维 Normal-Inverse-Wishart（NIW）证据协方差方法集成到同一"
        "轻量化网页界面中，用于辅助研究人员判断模型预测是否需要人工复核。",
    )
    add_body(
        document,
        "V1.0 同时实现模型直接推理与结果分析：DFUN 可从粉末衍射曲线执行 MC Dropout 空间群分类，"
        "EviField 可从两张 RGB 图像预测二维光流及 NIW 证据参数；用户也可上传已有推理结果。系统完成"
        "数值校验、指标计算、风险排序、图形展示和 CSV 导出，并提供账户、项目与操作审计。软件不承担"
        "神经网络在线训练或实验设备控制。",
    )
    document.add_page_break()
    document.add_heading("1.2 系统架构", level=2)
    add_body(
        document,
        "系统采用分层式轻量架构。浏览器端负责接收输入文件、设置推理参数和展示分析结果；Streamlit"
        "服务负责页面路由、会话状态和业务调用；模型适配层负责加载论文模型与参数文件；分析层使用"
        "NumPy、pandas 完成不确定性统计；SQLite 保存账户、项目及操作审计记录。",
    )
    add_table(
        document,
        ["层次", "组成", "职责"],
        [
            ["交互层", "Streamlit", "注册登录、项目管理、文件上传、结果展示、风险筛选与下载"],
            ["分析层", "分类 UQ / 2D NIW / 风险评分", "计算预测熵、互信息、协方差分解、几何诊断和风险分数"],
            ["模型层", "DFUN / EviField 适配器", "延迟加载论文模型、预处理输入并生成规范化分析表"],
            ["数据层", "SQLite / CSV / 内存数据表", "保存账户、项目和审计日志；接收推理结果并生成分析表"],
        ],
        [2.2, 4.2, 9.0],
    )
    add_body(document, "该结构将模型推理与不确定性分析解耦。未安装模型依赖时，用户仍可对已有推理结果执行全部核心分析。")
    add_body(
        document,
        "DFUN 和 EviField 均通过统一适配器输出标准数据表。分类适配器输出 sample_id、pass_id、"
        "class_label、probability 等字段；矢量场适配器输出 x、y、mean_1、mean_2、kappa、nu 及"
        "Cholesky 参数。界面与分析内核不依赖原模型内部张量命名，便于后续替换参数版本。",
    )

    document.add_page_break()
    document.add_heading("1.3 软件运行环境", level=2)
    add_table(
        document,
        ["组件", "最低要求", "推荐配置"],
        [
            ["操作系统", "Windows 10 / macOS 12 / Ubuntu 20.04", "64 位桌面或服务器系统"],
            ["Python", "3.10", "3.11 或 3.12"],
            ["内存", "8 GB", "16 GB 以上"],
            ["磁盘", "2 GB 可用空间", "10 GB 以上（不含模型参数）"],
            ["浏览器", "Chrome / Edge / Firefox 最新版", "分辨率 1280×720 以上"],
            ["GPU", "非必需，可使用 CPU", "Apple MPS、CUDA 或其他 PyTorch 加速设备"],
        ],
        [3.0, 5.6, 6.8],
    )
    add_body(
        document,
        "客户端只需要现代浏览器，无需单独安装插件。结果分析模式不要求 GPU；模型直接推理可以在 CPU"
        "上运行，也会自动检测 CUDA 或 Apple MPS。EviField 在较高分辨率下对内存要求较高，建议先用"
        "96×64 快速模式确认输入无误，再根据实际精度需求提高分辨率。",
    )
    document.add_heading("1.4 软件功能概述", level=2)
    add_bullets(
        document,
        [
            "用户与权限管理：支持注册、登录、退出和密码修改；首个账户为管理员，后续账户为分析用户。",
            "项目与审计管理：按任务建立分析项目，管理员可管理用户状态并查看操作日志。",
            "分类不确定性分析：汇总同一样本的多次类别概率，计算预测类别、Top-k、预测熵、互信息、概率方差和采样分歧。",
            "证据协方差分析：从二维 NIW 参数计算偶然、认知及总预测协方差，并生成迹、特征值、各向异性、相关性与椭圆参数。",
            "统一风险分级：将不同任务的不确定性指标组合为 0 至 100 的排序分数，并划分低、中、高和严重风险。",
            "风险覆盖分析：在提供真实标签或真实矢量时，计算保留不同覆盖率时的错误风险。",
            "结果筛选与导出：按风险分数筛选需复核对象，并将完整结果或高风险清单导出为 CSV。",
            "DFUN 模型推理：读取 NPZ 衍射曲线，提取 45 维峰值特征并执行多次随机前向传播。",
            "EviField 模型推理：读取双帧 RGB 图像，预测二维光流和像素级 NIW 证据参数。",
            "预训练模型管理：展示两个论文模型的权重路径、任务类型、输出协议和接入状态。",
        ],
    )
    document.add_heading("1.5 用户角色说明", level=2)
    add_body(document, "系统包含系统管理员和分析用户两类角色。角色由账户注册顺序和管理员控制规则确定。")
    add_table(
        document,
        ["角色", "可见功能", "权限范围"],
        [
            ["系统管理员", "全部分析功能、用户管理、操作日志", "查看全部项目，启用或停用普通账户，查看系统审计记录"],
            ["分析用户", "项目空间、两类分析、风险筛选、模型管理、账户设置", "仅查看和维护本人项目，不能管理其他账户"],
        ],
        [3.0, 5.5, 6.9],
    )
    add_body(
        document,
        "首次注册的账户自动成为系统管理员，后续注册账户默认为分析用户。管理员不能停用当前登录账户；"
        "普通用户被停用后不能继续登录，但已保存的项目和审计记录仍保留在数据库中。",
    )

    page_break(document)
    document.add_heading("二、软件使用说明", level=1)
    document.add_heading("2.1 启动软件", level=2)
    add_steps(
        document,
        [
            "进入软件根目录并创建 Python 虚拟环境。",
            "执行 pip install -e . 安装结果分析依赖；需要直接运行模型时执行 pip install -e '.[models]'。",
            "执行 streamlit run app.py 启动本地服务。",
            "在浏览器打开命令行显示的本地地址，进入账户登录页面。",
        ],
    )
    document.add_heading("2.1.1 注册与登录", level=3)
    add_body(
        document,
        "首次运行时切换到“创建账户”，填写显示名称、用户名和密码。首个注册账户自动成为系统管理员，"
        "后续注册账户默认为分析用户。密码至少八位并同时包含字母和数字，数据库只保存 scrypt 加盐哈希。",
    )
    add_screenshot(document, "login-v2.png", "图2.1 用户登录页面")
    document.add_heading("2.1.2 退出登录", level=3)
    add_body(
        document,
        "点击左侧导航栏底部的“退出登录”按钮，系统写入退出审计记录、清理当前浏览器会话并返回登录页。"
        "为防止分析结果被其他人员继续访问，在公共计算机或共享工作站使用完毕后应主动退出。",
    )
    document.add_heading("2.1.3 系统总览", level=3)
    add_screenshot(document, "dashboard-v3.png", "图2.2 系统总览页面")
    add_body(document, "系统总览展示项目数、有效用户、个人操作记录、模型状态和核心分析能力。管理员侧栏额外显示用户管理与操作日志。")

    document.add_heading("2.2 项目与用户基础管理", level=2)
    add_body(
        document,
        "分析用户可在项目空间新建项目，记录项目名称、编号、任务类型和说明，并仅查看自己的项目。系统"
        "管理员可查看全部项目，在用户管理页面启用或停用普通账户，并通过操作日志审计注册、登录、项目"
        "变更、分析执行和密码修改等操作。管理员不能停用当前登录账户。",
    )
    add_table(
        document,
        ["项目字段", "填写要求", "用途"],
        [
            ["项目名称", "不能为空，建议体现材料或任务名称", "用于项目卡片和列表展示"],
            ["项目编号", "同一用户下建议保持唯一", "用于样本批次、实验编号或外部任务关联"],
            ["任务类型", "分类不确定性或二维矢量场", "区分项目对应的分析流程"],
            ["项目说明", "可填写数据来源、模型版本和复核目的", "为后续审计和结果解释保留上下文"],
        ],
        [3.0, 6.0, 6.4],
    )
    add_steps(
        document,
        [
            "进入“项目空间”，展开“新建分析项目”。",
            "填写项目名称、项目编号、任务类型和项目说明后保存。",
            "在项目列表中确认项目状态、创建时间和任务类型。",
            "如需删除项目，展开删除区域并二次确认；删除后项目记录不可恢复。",
        ],
    )
    add_screenshot(document, "project-v3.png", "图2.3 项目空间页面")

    page_break(document)
    document.add_heading("2.3 分类不确定性分析", level=2)
    document.add_heading("2.3.1 输入文件", level=3)
    add_body(
        document,
        "分类页面提供“DFUN模型直接推理”和“导入概率结果”两种输入方式。前者使用训练好的空间群"
        "分类模型从衍射曲线生成多次类别概率，后者用于分析其他程序已经生成的 Monte Carlo 预测结果。",
    )
    add_table(
        document,
        ["NPZ键名", "数据形式", "说明"],
        [
            ["intensity", "一维数组或 N×L 数组", "原始衍射强度；二维时通过样本序号选择一行"],
            ["features", "一维数组或 N×L 数组", "与 intensity 二选一，用于兼容模型包公开样本"],
            ["d_spacing", "长度为 L 的一维数组", "非5000点曲线必须提供，用于重采样到训练网格"],
            ["labels230", "整数数组", "可选，零起始类别编号；系统转换为空间群1至230"],
            ["space_group", "整数", "可选，直接记录真实空间群编号"],
        ],
        [3.2, 5.0, 7.2],
    )
    add_body(
        document,
        "分类模块读取长表 CSV。必要字段为 sample_id、pass_id、class_label 和 probability，可选字段 true_label。"
        "同一样本必须具有相同采样次数，每次采样必须包含相同类别集合。系统允许概率行和不等于 1，并在验证通过后自动归一化。",
    )
    add_table(
        document,
        ["字段", "含义", "约束"],
        [
            ["sample_id", "样本唯一编号", "同一样本可出现多行"],
            ["pass_id", "随机前向传播轮次", "各样本轮次数一致"],
            ["class_label", "类别名称或编号", "每轮类别集合一致"],
            ["probability", "当前类别概率", "有限且非负"],
            ["true_label", "真实类别", "可选，用于准确率、ECE 和风险曲线"],
        ],
        [3.0, 6.0, 6.4],
    )
    document.add_heading("2.3.2 执行分析", level=3)
    add_steps(
        document,
        [
            "选择“DFUN模型直接推理”，上传衍射曲线 NPZ 或使用内置公开样本，设置 MC Dropout 次数后运行模型。",
            "如已有模型结果，切换至“导入概率结果”，上传符合格式的 CSV 或载入演示数据。",
            "展开输入预览检查字段和值域，系统自动或按按钮执行不确定性分析。",
            "按 risk_score 降序检查高风险样本，必要时下载完整结果。",
        ],
    )
    add_body(
        document,
        "MC Dropout 次数允许设置为5至100次，默认30次。次数越多，概率均值和方差估计通常越稳定，"
        "但推理耗时近似按次数线性增加。计算设备选择“自动选择”时，系统依次检测 CUDA、Apple MPS，"
        "均不可用时回退到 CPU。样本序号用于从包含多条曲线的 NPZ 中选择待推理样本。",
    )
    add_body(
        document,
        "DFUN 推理过程中模型整体保持评估模式，仅将 Dropout 层切换到随机状态，BatchNorm 层不会更新"
        "统计量。每次前向传播产生230个类别概率，系统将空间群编号、概率、推理轮次和门控权重整理成"
        "长表，再调用统一分类不确定性分析内核。",
    )

    page_break(document)
    document.add_heading("2.3.3 分类结果说明", level=3)
    add_screenshot(document, "classification-result-v3.png", "图2.4 分类不确定性分析结果页面")
    add_table(
        document,
        ["指标", "解释"],
        [
            ["confidence", "平均概率中最大值，越大表示最终类别更集中"],
            ["predictive_entropy", "平均预测分布的熵，越大表示类别整体越模糊"],
            ["normalized_entropy", "预测熵除以类别数对数，用于不同类别数任务间比较"],
            ["expected_entropy", "每次随机前向预测熵的平均值"],
            ["mutual_information", "预测熵减去采样期望熵，用于描述多次随机推理之间的分歧"],
            ["variation_ratio", "多次推理的获胜类别不一致比例"],
            ["mean_probability_variance", "各类别在多次随机推理中的概率方差均值"],
            ["top1至top3", "平均概率排序前三的类别及对应概率"],
            ["risk_score", "0 至 100 的综合排序分数，不解释为真实错误概率"],
        ],
        [5.0, 10.4],
    )

    page_break(document)
    document.add_heading("2.4 二维矢量场证据协方差分析", level=2)
    document.add_heading("2.4.1 输入文件", level=3)
    add_body(
        document,
        "矢量场页面提供“EviField模型直接推理”和“导入NIW结果”两种方式。模型直接推理需要按顺序"
        "提供参考帧 Frame 1 与支持帧 Frame 2，两张图像应具有相同尺寸，支持 PNG、JPG、JPEG、TIF"
        "和 TIFF。系统将图像转换为 RGB 浮点张量，并由0至255或0至1归一化到[-1,1]。",
    )
    add_table(
        document,
        ["推理分辨率", "输出像素数", "适用场景"],
        [
            ["96×64（快速）", "6144", "输入检查、功能演示和CPU快速推理"],
            ["192×128", "24576", "一般预览和中等精度分析"],
            ["384×256", "98304", "较精细空间结构分析，需要更多内存"],
            ["保持原始尺寸", "由输入图像决定", "最终分析或需要保留原始坐标时使用"],
        ],
        [4.0, 4.0, 7.4],
    )
    add_body(
        document,
        "二维矢量场模块读取像素级 NIW 参数 CSV。必要字段为 x、y、mean_1、mean_2、kappa、nu、l11、l21、l22。"
        "其中 l11、l21、l22 构成尺度矩阵的下三角 Cholesky 因子；kappa 必须大于 0，nu 必须大于 3，"
        "l11 与 l22 必须大于 0。可选 target_1 和 target_2 用于端点误差与风险覆盖分析。",
    )
    add_steps(
        document,
        [
            "选择“EviField模型直接推理”，上传参考帧与支持帧，选择分辨率后运行模型；也可使用内置双帧图像。",
            "如已有证据结果，切换至“导入NIW结果”，上传像素级 CSV 或载入演示矢量场。",
            "模型或导入数据完成后，系统计算协方差分解和几何诊断。",
            "在热力图下拉框中切换统一风险、协方差迹、各向异性、相关性和证据风险。",
            "检查像素级结果并导出 CSV。",
        ],
    )
    add_body(document, "软件以输入坐标 x、y 重构二维热力图。若坐标不构成规则网格，表格结果仍可正常计算，但热力图可能包含空白位置。")
    add_body(
        document,
        "若外部系统对图像进行了缩放，模型预测的位移分量位于缩放后坐标系中。需要回到原图坐标时，"
        "应在恢复空间尺寸的同时分别按宽度和高度比例缩放水平、垂直分量。当前软件用于不确定性诊断，"
        "不会自动修改用户上传的原始文件。",
    )

    page_break(document)
    document.add_heading("2.4.2 矢量场结果说明", level=3)
    add_screenshot(document, "vector-result-v3.png", "图2.5 二维矢量场证据协方差分析结果页面")
    add_table(
        document,
        ["指标", "解释"],
        [
            ["trace_uncertainty", "总预测协方差迹的一半，描述局部总体不确定性尺度"],
            ["aleatoric_11/12/22", "偶然协方差的三个独立元素"],
            ["epistemic_11/12/22", "认知协方差的三个独立元素"],
            ["total_11/12/22", "总预测协方差的三个独立元素"],
            ["anisotropy_ratio", "最大与最小特征值之比，描述方向不均衡程度"],
            ["correlation", "两个矢量分量预测误差的归一化相关性"],
            ["principal_angle_deg", "最大特征值对应的主不确定方向"],
            ["ellipse_major/minor_95", "95% 椭圆的长轴和短轴直径"],
            ["evidence_risk", "由 kappa 与 nu 得到的证据不足分数，值越大风险越高"],
            ["endpoint_error", "提供真实矢量时，预测均值与真实矢量的欧氏距离"],
        ],
        [5.0, 10.4],
    )

    page_break(document)
    document.add_heading("2.5 高风险筛选与导出", level=2)
    add_body(
        document,
        "完成任一分析后进入“高风险筛选”页面。风险阈值滑块默认设置为 60，页面仅显示不低于该阈值的"
        "样本或像素。列表包含任务类型、对象编号、风险分数、风险等级、主要风险原因和复核建议。",
    )
    add_bullets(
        document,
        [
            "低风险：0 ≤ 分数 < 30；",
            "中风险：30 ≤ 分数 < 60；",
            "高风险：60 ≤ 分数 < 80；",
            "严重风险：80 ≤ 分数 ≤ 100。",
        ],
    )
    add_body(
        document,
        "风险分数用于排序、筛选和分配人工复核资源，不是模型预测错误概率。没有真实标签时，软件只输出"
        "相对风险；存在真实标签时，Risk-Coverage 曲线用于检验风险排序是否能优先剔除错误对象。",
    )
    add_screenshot(document, "risk-v3.png", "图2.6 高风险结果筛选页面")
    add_body(
        document,
        "风险阈值可以在0至100之间调整，任务类型可选择全部、分类或矢量场。筛选结果按风险分数降序"
        "排列，并保留风险原因和复核建议。点击“导出高风险清单”可下载带有UTF-8 BOM的CSV文件，"
        "便于在常用表格软件中直接打开中文内容。",
    )
    document.add_heading("2.6 模型管理", level=2)
    add_body(
        document,
        "模型管理页面检查 DFUN 与 EviField 模型包的网络代码和参数文件。当前两个模型均显示“已接入”。"
        "inference/adapters.py 采用延迟加载方式：只有用户执行直接推理时才导入 PyTorch 和模型结构，"
        "并将输出转换为分类长表或像素级 NIW 标准表。",
    )
    add_table(
        document,
        ["模型", "任务", "标准输出", "参数状态"],
        [
            ["DFUN", "230类空间群分类", "T×230类别概率", "已接入"],
            ["EviField", "双帧二维光流", "均值、κ、ν、L及全协方差", "已接入"],
        ],
        [3.0, 4.0, 6.0, 2.4],
    )
    add_screenshot(document, "model-v3.png", "图2.7 预训练模型管理页面")

    document.add_heading("2.7 账户设置", level=2)
    add_body(
        document,
        "账户设置页面显示当前用户的显示名称、用户名、角色、状态和创建时间。用户可输入当前密码、"
        "新密码和确认密码完成修改。新密码至少八位并同时包含字母和数字；当前密码错误或两次新密码"
        "不一致时，系统拒绝修改并显示原因。",
    )
    add_steps(
        document,
        [
            "进入“账户设置”，核对当前登录账户和角色。",
            "在“当前密码”中输入原密码。",
            "输入并确认符合规则的新密码。",
            "点击“更新密码”；修改成功后使用新密码完成后续登录。",
        ],
    )
    add_screenshot(document, "account-v3.png", "图2.8 账户设置页面")

    document.add_heading("2.8 用户管理", level=2)
    add_body(
        document,
        "用户管理页面仅管理员可见。页面展示账户总数、有效账户数、管理员数量和账户卡片。管理员可"
        "选择普通账户并将状态设置为正常或已停用；系统禁止管理员停用自身账户。状态变更写入审计日志。",
    )
    add_table(
        document,
        ["字段", "说明"],
        [
            ["用户名", "登录使用的唯一名称，匹配时不区分大小写重复"],
            ["显示名称", "侧栏和账户卡片显示的名称"],
            ["角色", "系统管理员或分析用户"],
            ["状态", "正常账户可以登录；已停用账户不能登录"],
            ["创建时间", "账户首次注册并写入数据库的时间"],
        ],
        [5.0, 10.4],
    )
    add_screenshot(document, "users-v3.png", "图2.9 用户管理页面")

    document.add_heading("2.9 操作日志", level=2)
    add_body(
        document,
        "操作日志用于审计注册、登录、退出、密码修改、项目创建与删除、模型推理、结果分析和账户状态"
        "变更。管理员可以按操作类型和操作用户筛选记录。日志详情采用结构化字段保存模型、行数、像素数、"
        "推理次数、设备类型等必要信息，不保存用户密码和上传文件的原始内容。",
    )
    add_screenshot(document, "logs-v3.png", "图2.10 操作日志页面")

    page_break(document)
    document.add_heading("三、核心算法说明", level=1)
    document.add_heading("3.1 DFUN 输入预处理", level=2)
    add_body(
        document,
        "DFUN训练网格为 np.flip(np.linspace(0.889,17.659,5000))。当输入曲线自带d_spacing时，系统先"
        "按位置升序执行线性插值，再恢复训练网格的下降方向；当输入没有d_spacing时，强度长度必须"
        "等于5000。NaN和无穷值转换为0，负强度截断为0，随后按最大值归一化到0至100。",
    )
    add_body(
        document,
        "峰值分支使用 scipy.signal.find_peaks 检测显著性不低于最大强度5%的峰，按显著性选取前10个。"
        "每个峰提取d位置、峰强、d间隔单位的半高宽和不对称性4项特征；不足10个峰时补0。最后增加"
        "均值、标准差、偏度、峰度和强度质心5项全局统计量，形成45维物理特征。",
    )
    document.add_heading("3.2 Monte Carlo Dropout 分类量化", level=2)
    add_body(
        document,
        "对每个样本执行 T 次启用 Dropout 的随机前向传播，得到类别概率 p_t。软件先计算平均概率 p̄，"
        "再取最大概率对应类别为最终预测。预测熵为 H(p̄) = -Σ p̄_k ln(p̄_k)。采样期望熵为各次预测熵"
        "的平均值，互信息为 MI = H(p̄) - E[H(p_t)]。该过程对应 DFUN 论文采用 T=50 次前向传播评估"
        "预测均值、方差与熵的思路。",
    )
    add_body(
        document,
        "分类风险分数由归一化预测熵、1-confidence、平均概率方差和 variation ratio 按固定权重组合："
        "Risk_cls = 100 × (0.40H_norm + 0.30(1-c) + 0.15V_norm + 0.15R_var)。所有分量限制在 0 至 1。",
    )
    document.add_heading("3.3 分类风险评分", level=2)
    add_body(
        document,
        "系统将预测熵除以ln(K)得到归一化熵H_norm；将平均概率方差除以理论上限0.25并截断到0至1；"
        "variation ratio定义为1减去最终获胜类别在各次随机推理中获胜的频率。分类风险分数采用固定"
        "权重组合，使熵、置信度、方差和获胜类别分歧同时参与排序。",
    )
    document.add_heading("3.4 风险覆盖", level=2)
    add_body(
        document,
        "软件按不确定性从低到高排序对象，依次增加保留数量。Coverage 为已保留对象占全部对象的比例；"
        "Risk 为保留对象的累计平均错误。理想风险排序应使低覆盖率区域保持较低风险，并随覆盖率上升而"
        "逐步接近全体数据风险。",
    )

    page_break(document)
    document.add_heading("3.5 二维 NIW 协方差分解", level=2)
    add_body(
        document,
        "EviField 在每个像素输出均值向量 m、证据参数 κ、自由度 ν 以及正定尺度矩阵 Ψ。软件通过"
        "L=[[l11,0],[l21,l22]] 与 Ψ=LLᵀ+εI 保证尺度矩阵正定。二维条件下 ν>3。",
    )
    add_bullets(
        document,
        [
            "偶然协方差：C_ale = Ψ / (ν - 3)；",
            "认知协方差：C_epi = Ψ / [κ(ν - 3)]；",
            "总预测协方差：C_tot = (κ + 1)Ψ / [κ(ν - 3)]；",
            "证据风险：s_E = -ln(1 + κ) - ln(1 + ν - 3 + ε)。",
        ],
    )
    add_body(
        document,
        "对 C_tot 进行对称特征分解得到最大、最小特征值与主特征向量。软件进一步计算迹、特征值间隙、"
        "各向异性比、通道相关系数、条件数、主方向角和 95% 椭圆轴长。NIW 所给出的偶然/认知分解应"
        "理解为模型内部诊断量，不等同于严格可证明的因果不确定性分解。",
    )
    document.add_heading("3.6 矢量场风险评分", level=2)
    add_body(
        document,
        "矢量场风险由证据风险、总协方差迹、最大特征值和各向异性四个分量组成。前三个分量先在当前"
        "结果内部转换为百分位分数，各向异性使用特征值间隙归一化量。最终权重依次为0.45、0.30、"
        "0.15和0.10，并映射到0至100。该分数用于当前批次的相对排序，不应跨数据域直接比较。",
    )
    add_body(
        document,
        "风险原因按优先级生成：证据分量达到高百分位时标记“模型证据不足”；总协方差迹达到高百分位"
        "时标记“总协方差较大”；各向异性较高时标记“局部各向异性显著”；其余位置标记为综合风险较低。",
    )

    page_break(document)
    document.add_heading("四、数据校验与异常处理", level=1)
    document.add_heading("4.1 分类数据校验", level=2)
    add_bullets(
        document,
        [
            "缺少必要字段时终止分析并列出缺失字段；",
            "概率为负数、空值或无穷值时拒绝输入；",
            "同一样本、轮次和类别出现重复记录时拒绝输入；",
            "各轮类别集合不一致或样本采样次数不一致时拒绝输入；",
            "每轮概率总和大于 0 时自动归一化。",
        ],
    )
    document.add_heading("4.2 NIW 数据校验", level=2)
    add_bullets(
        document,
        [
            "全部参数必须为有限数值；",
            "κ≤0、ν≤3 或 Cholesky 对角元素非正时拒绝输入；",
            "可选真实矢量必须同时包含 target_1 与 target_2；",
            "特征分解使用对称矩阵算法，并添加 εI 提高数值稳定性。",
        ],
    )
    document.add_heading("4.3 模型输入与加载异常", level=2)
    add_bullets(
        document,
        [
            "DFUN输入缺少intensity和features时拒绝推理；非5000点曲线缺少d_spacing时给出提示；",
            "DFUN的MC Dropout次数小于2时拒绝执行，避免将单次确定性输出误认为随机采样；",
            "EviField必须同时提供两张同尺寸RGB图像，单帧、尺寸不一致或通道异常时拒绝输入；",
            "模型参数文件不存在时显示缺失路径；权重键不匹配时列出missing和unexpected项目；",
            "缺少PyTorch、SciPy或Pillow时提示执行pip install -e '.[models]'，基础结果分析仍可使用；",
            "GPU或MPS不可用时自动回退CPU；设备执行异常会显示错误，不生成虚假分析结果。",
        ],
    )
    document.add_heading("4.4 账户与数据安全", level=2)
    add_body(
        document,
        "密码在写入数据库前使用scrypt算法和随机盐计算哈希，不保存明文。账户状态、项目归属和管理员"
        "权限均在服务端校验。上传的CSV、NPZ和图像在当前会话中处理，不写入SQLite；用户主动导出的"
        "结果由浏览器下载到本地。模型参数文件不应提交到公开代码仓库。",
    )
    document.add_heading("4.5 测试结果", level=2)
    add_body(
        document,
        "当前版本包含10项自动化单元测试，覆盖分类确定性样本、字段缺失、NIW协方差公式、自由度约束、"
        "注册登录、密码修改、账户停用、项目归属、审计日志、模型参数发现和EviField标准输出转换。测试"
        "同时验证分类互信息在完全一致的重复预测下为零，以及转换后的二维NIW总协方差与模型包预期"
        "C_total输出一致。",
    )
    add_body(
        document,
        "真实权重冒烟测试中，DFUN公开样本经5次MC Dropout后预测空间群148，与模型包预期类别一致；"
        "EviField公开双帧样本输出6144个像素，kappa和nu均值以及总不确定性与预期摘要在浮点误差范围内"
        "一致。浏览器端还验证了30次DFUN推理、EviField快速分辨率推理、风险结果展示和模型状态页面。",
    )

    page_break(document)
    document.add_heading("五、预训练模型实现说明", level=1)
    document.add_heading("5.1 DFUN 模型", level=2)
    add_body(
        document,
        "DFUN 使用 dfun_gate_fls_exp_test4_model.pth 参数文件。输入曲线重采样到 5000 点并归一化至 0-100，"
        "同时提取 10 个主要峰的峰位、强度、半高宽和不对称性及 5 个全局统计量，共 45 维物理特征。模型"
        "输出 230 个空间群概率；推理时保持 BatchNorm 为评估状态，仅启用 Dropout，生成 sample_id、"
        "pass_id、class_label 与 probability 长表。",
    )
    document.add_heading("5.2 EviField 模型", level=2)
    add_body(
        document,
        "EviField 使用 evifield_era_optical_flow_latest_S_Mix.pt 参数文件。输入为参考帧与支持帧两张 RGB"
        "图像，像素归一化到 [-1,1]。网络输出二维光流均值、κ、ν 和下三角矩阵 L。适配器依据"
        "Ψ=νLLᵀ 将模型内部参数转换为分析内核所需的 Cholesky 参数，再执行偶然、认知和总协方差诊断。",
    )
    document.add_heading("5.3 运行与版本边界", level=2)
    add_body(
        document,
        "基础结果分析依赖 NumPy 与 pandas，可在 CPU 环境运行。模型直接推理额外依赖 PyTorch、SciPy"
        "和 Pillow，支持 CPU、Apple MPS 或 CUDA。两个适配器均校验权重路径和参数键，加载失败时只显示"
        "错误信息，不生成伪造结果。当前版本不提供训练、微调和模型参数修改功能。",
    )
    document.add_heading("5.4 模型文件与代码目录", level=2)
    add_table(
        document,
        ["路径", "内容"],
        [
            ["dfun_model_package/checkpoints/", "DFUN训练参数文件"],
            ["dfun_model_package/dfun/", "DFUN网络、预处理和参考推理代码"],
            ["dfun_model_package/mappings/", "230类空间群编号映射"],
            ["system_handoff_evifield_era/checkpoints/", "EviField光流与LTEM参数文件"],
            ["system_handoff_evifield_era/code/", "EviField网络、NIW变换和预处理代码"],
            ["inference/adapters.py", "统一的延迟加载和标准表转换接口"],
            ["uqcore/", "分类不确定性、NIW协方差和风险评分内核"],
        ],
        [8.0, 7.4],
    )
    document.add_heading("5.5 软件部署与启动", level=2)
    add_steps(
        document,
        [
            "确认Python版本不低于3.10，并在软件根目录创建虚拟环境。",
            "激活虚拟环境，执行pip install -e '.[models]'安装全部运行依赖。",
            "执行streamlit run app.py启动服务；默认数据库位于storage/sciuq.db。",
            "需要指定独立数据库时，设置SCIUQ_DB环境变量后再启动。",
            "首次访问时创建管理员账户，并在模型管理页确认两个模型均显示“已接入”。",
            "依次使用内置XRD样本和双帧图像执行一次推理，确认设备、权重和输出链路正常。",
        ],
    )
    add_body(
        document,
        "生产或实验室共享部署时，建议由操作系统服务管理Streamlit进程，并限制数据库与模型参数目录"
        "的文件权限。外部访问场景应在前端配置HTTPS反向代理。软件本身不上传数据到云端，网络暴露"
        "范围由部署方的服务器、防火墙和代理配置决定。",
    )

    OUTPUT.mkdir(parents=True, exist_ok=True)
    path = OUTPUT / "产品文档-异构材料数据不确定性分析系统-V1.0.docx"
    document.save(path)
    return path


def replace_cell_text(cell, text: str) -> None:
    first_paragraph = cell.paragraphs[0]
    if first_paragraph.runs:
        first_paragraph.runs[0].text = text
        for run in first_paragraph.runs[1:]:
            run.text = ""
    else:
        first_paragraph.add_run(text)
    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def build_application() -> Path:
    document = Document(APPLICATION_TEMPLATE)
    table = document.tables[0]
    replacements = {
        (1, 4): "SciUQ Studio",
        (14, 8): "PC机，4核CPU，16G内存，20G可用硬盘",
        (15, 8): "PC或服务器，4核CPU，8G内存，10G硬盘，GPU可选",
        (16, 11): "macOS 12及以上或Ubuntu 20.04",
        (17, 11): "Python3.10+，Streamlit，SQLite，NumPy，pandas，Plotly",
        (18, 11): "Windows 10、macOS 12或Ubuntu 20.04及以上",
        (19, 11): "Python3.10+，Streamlit，SQLite3，PyTorch，SciPy，现代浏览器",
        (20, 4): "Python3.10及以上",
        (20, 12): f"{core_source_line_count()}行",
        (23, 9): (
            "支持用户注册登录、项目管理和日志审计；运行DFUN空间群分类和EviField光流模型，或导入已有"
            "概率与二维NIW参数，计算预测熵、协方差分解及证据风险，实现风险分级、筛选、可视化和导出。"
        ),
        (25, 6): (
            "采用Streamlit和SQLite轻量架构，scrypt加盐保护密码，NumPy向量化计算、Plotly展示风险图表；"
            "通过PyTorch延迟加载DFUN和EviField权重，模型推理与不确定性分析内核解耦。"
        ),
    }
    for (row_index, cell_index), value in replacements.items():
        replace_cell_text(table.rows[row_index].cells[cell_index], value)
    # This template cell contains a highlighted hyperlink-like fragment that is
    # not exposed as a normal python-docx run. Rebuild the cell to remove it.
    table.rows[25].cells[6].text = replacements[(25, 6)]
    OUTPUT.mkdir(parents=True, exist_ok=True)
    path = OUTPUT / "软著申请表-异构材料数据不确定性分析系统-V1.0-草稿.docx"
    document.save(path)
    return path


if __name__ == "__main__":
    print(build_manual())
    print(build_application())
